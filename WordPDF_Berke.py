import docx
import PyPDF2
import re
import difflib
from collections import Counter
from docx import Document
from docx.shared import RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

class DocumentTextFormatter:
    def __init__(self, file_path, file_type):
        self.file_path = file_path
        self.file_type = file_type
        self.text = ""

    def extract_text_from_docx(self):
        doc = docx.Document(self.file_path)
        for para in doc.paragraphs:
            self.text += para.text + "\n"

    def extract_text_from_pdf(self):
        with open(self.file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                self.text += reader.pages[page_num].extract_text() + "\n"

    def extract_text(self):
        if self.file_type == 'docx':
            self.extract_text_from_docx()
        elif self.file_type == 'pdf':
            self.extract_text_from_pdf()

    def format_text(self):
        self.text = re.sub(r'(\w)(\s{2,})(\w)', r'\1 \3', self.text)
        self.text = re.sub(r'\s+', ' ', self.text)
        self.text = re.sub(r'\s([.,;:!?])', r'\1', self.text)
        self.text = re.sub(r'([.,;:!?])([A-Za-z])', r'\1 \2', self.text)
        self.text = re.sub(r'(\d)(\s+)(\d)', r'\1\3', self.text)
        self.text = re.sub(r'\s+\n', '\n', self.text).strip()
        
        formatted_text = ""
        lines = self.text.split('\n')
        for line in lines:
            if line.strip():
                formatted_text += line.strip() + ' '
            else:
                formatted_text = formatted_text.strip() + '\n\n'
        
        self.text = re.sub(r'\n\n+', '\n\n', formatted_text)

    def get_formatted_text(self):
        self.extract_text()
        self.format_text()
        return self.text.strip()

def highlight_differences(text1, text2):
    diff = difflib.ndiff(text1.split(), text2.split())
    highlighted_text1 = []
    highlighted_text2 = []

    for word in diff:
        if word.startswith('- '):
            highlighted_text1.append(('red', word[2:]))
        elif word.startswith('+ '):
            highlighted_text2.append(('green', word[2:]))
        else:
            word = word[2:]
            highlighted_text1.append(('black', word))
            highlighted_text2.append(('black', word))
    
    return highlighted_text1, highlighted_text2


def compare_docs(file_path1, file_type1, file_path2, file_type2):
    formatter1 = DocumentTextFormatter(file_path1, file_type1)
    text1 = formatter1.get_formatted_text()
    
    formatter2 = DocumentTextFormatter(file_path2, file_type2)
    text2 = formatter2.get_formatted_text()

    sequence_matcher = difflib.SequenceMatcher(None, text1, text2)
    similarity_percentage = sequence_matcher.ratio() * 100

    words1 = Counter(re.findall(r'\w+', text1))
    words2 = Counter(re.findall(r'\w+', text2))
    
    diff_words1 = words1 - words2
    diff_words2 = words2 - words1
    
    different_words_count = sum(diff_words1.values()) + sum(diff_words2.values())

    highlighted_text1, highlighted_text2 = highlight_differences(text1, text2)

    return similarity_percentage, different_words_count, highlighted_text1, highlighted_text2


def create_word(highlighted_text, output_path):
    doc = Document()
    para = doc.add_paragraph()
    
    for color, word in highlighted_text:
        run = para.add_run(word + ' ')
        if color == 'red':
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif color == 'green':
            run.font.color.rgb = RGBColor(0, 255, 0)

    doc.save(output_path)


def create_pdf(highlighted_text, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    story = []
    
    for color, word in highlighted_text:
        if color == 'red':
            style = ParagraphStyle(name='Red', parent=normal_style, textColor='red')
        elif color == 'green':
            style = ParagraphStyle(name='Green', parent=normal_style, textColor='green')
        else:
            style = normal_style
        
        story.append(Paragraph(word, style))
        story.append(Spacer(1, 0.2 * inch))  # To add space between words
    
    doc.build(story)

# Dosya yolları ve türleri
file_path1 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\Belge1.docx"
file_type1 = 'docx'
file_path2 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\1. Belge.pdf"
file_type2 = 'pdf'

# Dosyaları kıyaslama
similarity_percentage, different_words_count, highlighted_text1, highlighted_text2 = compare_docs(file_path1, file_type1, file_path2, file_type2)

# Benzerlik yüzdesini ve farklı kelime sayısını yazdırma
print(f"Benzerlik Yüzdesi: {similarity_percentage:.2f}%")
print(f"Farklı Kelime Sayısı: {different_words_count}")

# Terminalde farkları yazdırma
def print_highlighted_text(highlighted_text):
    for color, word in highlighted_text:
        if color == 'red':
            print(f"\033[91m{word}\033[0m", end=' ')
        elif color == 'green':
            print(f"\033[92m{word}\033[0m", end=' ')
        else:
            print(word, end=' ')
    print()

print("Birinci dosyanın farklı kelimeleri:")
print_highlighted_text(highlighted_text1)

print("İkinci dosyanın farklı kelimeleri:")
print_highlighted_text(highlighted_text2)

# Yeni dosyaları oluşturma
output_path1 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\Belge3.docx"
output_path2 = r"C:\Users\berkeb\OneDrive - skoda.com.tr\Masaüstü\4. Belge.pdf"

if file_type1 == 'docx':
    create_word(highlighted_text1, output_path1)
else:
    create_pdf(highlighted_text1, output_path1)

if file_type2 == 'pdf':
    create_pdf(highlighted_text2, output_path2)
else:
    create_word(highlighted_text2, output_path2)
